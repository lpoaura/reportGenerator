from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

from reportgenerator.analysis.common.tables.utils import (
    set_cell_background, set_cell_font, set_cell_text_color)
from reportgenerator.analysis.environmental_zones.summary_text import \
    TYPE_LABELS

LPO_BLUE = "0088CC"
LPO_WHITE = "FFFFFF"

ZONAGE_COLUMNS = [
    ("type_code", "Type de zonage"),
    ("zone_etude", "Zone d'étude\n(km²)"),
    ("buffer", "Périmètre autour\nde la zone d'étude\n(km²)"),
    ("10km", "10km autour\ndu périmètre\n(km²)"),
]

COLUMN_WIDTHS = {
    "type_code": Inches(2.5),
    "zone_etude": Inches(1.3),
    "buffer": Inches(1.3),
    "10km": Inches(1.3),
}


def insert_zonage_table(document, placeholder, data):
    for paragraph in document.paragraphs:

        if placeholder not in paragraph.text:
            continue

        paragraph.text = paragraph.text.replace(placeholder, "")

        table = document.add_table(rows=1, cols=len(ZONAGE_COLUMNS))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # HEADER
        header_cells = table.rows[0].cells
        for idx, (field, label) in enumerate(ZONAGE_COLUMNS):
            cell = header_cells[idx]
            cell.text = label
            cell.width = COLUMN_WIDTHS[field]
            set_cell_background(cell, LPO_BLUE)
            set_cell_text_color(cell, LPO_WHITE)
            set_cell_font(cell, bold=True, size=10)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # DATA
        for item in data:
            row_cells = table.add_row().cells
            for idx, (field, _) in enumerate(ZONAGE_COLUMNS):
                cell = row_cells[idx]
                if field == "type_code":
                    value = TYPE_LABELS.get(item["type_code"], item["type_code"])
                else:
                    value = f"{item.get(field, 0):.2f}"
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.text = str(value)
                cell.width = COLUMN_WIDTHS[field]
                set_cell_font(cell, size=10)

        paragraph._element.addnext(table._element)
        break