from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from analysis.common.tables.utils import (
    apply_status_style_species,
    italicize_cell_species,
    set_cell_background,
    set_cell_text_color,
    set_cell_font,
)

# =========================================================
# CONFIGURATION STANDARD DU TABLEAU ESPÈCES
# =========================================================

SPECIES_COLUMNS = [
    ("nom_vern", "Nom vernaculaire"),
    ("lb_nom", "Nom scientifique"),
    ("prot_nat", "Protection nationale"),
    ("lr_fr_nich", "LR France\nNicheur"),
    ("lr_fr_hiv", "LR France\nHivernant"),
    ("lr_fr_migr", "LR France\nMigration"),
    ("lr_aura", "LR AuRA"),
    ("nb_annees_observation", "Nb\nannées obs."),
    ("nb_observations", "Nb\ndonnées"),
    ("derniere_annee_observation", "Dernière\nobs."),
]

STATUS_FIELDS = [
    "lr_fr_nich",
    "lr_fr_hiv",
    "lr_fr_migr",
    "lr_aura",
]

SCIENTIFIC_FIELDS = [
    "lb_nom",
]

CENTER_FIELDS = [
    "lr_fr_nich",
    "lr_fr_hiv",
    "lr_fr_migr",
    "lr_aura",
    "nb_observations",
    "derniere_annee_observation",
    "nb_annees_observation",
]

# Largeurs des colonnes
COLUMN_WIDTHS = {
    "nom_vern": Inches(2.2),
    "lb_nom": Inches(2.3),
    "prot_nat": Inches(1.4),
    "lr_fr_nich": Inches(0.7),
    "lr_fr_hiv": Inches(0.9),
    "lr_fr_migr": Inches(0.9),
    "lr_aura": Inches(0.7),
    "nb_annees_observation": Inches(0.9),
    "nb_observations": Inches(1.0),
    "derniere_annee_observation": Inches(1.0),
}

# =========================================================
# COULEURS LPO
# =========================================================

LPO_BLUE = "0088CC"
LPO_WHITE = "FFFFFF"
ROW_BORDER = "D9D9D9"
LPO_GREY = "626a6e"

# =========================================================
# INSERTION TABLEAU
# =========================================================

def insert_species_table(
    document,
    placeholder,
    data
):

    for paragraph in document.paragraphs:

        if placeholder not in paragraph.text:
            continue

        paragraph.text = paragraph.text.replace(
            placeholder,
            ""
        )

        # =====================================================
        # CREATION TABLE
        # =====================================================

        table = document.add_table(
            rows=1,
            cols=len(SPECIES_COLUMNS)
        )

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # =====================================================
        # HEADER
        # =====================================================

        header_cells = table.rows[0].cells

        for idx, (field, label) in enumerate(SPECIES_COLUMNS):

            cell = header_cells[idx]

            cell.text = label

            # largeur
            cell.width = COLUMN_WIDTHS[field]

            # style fond
            set_cell_background(
                cell,
                LPO_BLUE
            )

            # texte
            set_cell_text_color(
                cell,
                LPO_WHITE
            )

            # police
            set_cell_font(
                cell,
                bold=True,
                size=10
            )

            # alignement
            paragraph_header = cell.paragraphs[0]
            paragraph_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # =====================================================
        # DONNÉES
        # =====================================================

        for item in data:

            row_cells = table.add_row().cells

            for idx, (field, _) in enumerate(SPECIES_COLUMNS):

                cell = row_cells[idx]

                value = item.get(field, "")

                # sécurité None
                if value is None:
                    value = ""

                cell.text = str(value)

                # largeur
                cell.width = COLUMN_WIDTHS[field]

                # police générale
                set_cell_font(
                    cell,
                    size=10
                )

                # centrage
                if field in CENTER_FIELDS:

                    cell.paragraphs[0].alignment = (
                        WD_PARAGRAPH_ALIGNMENT.CENTER
                    )

                # italique scientifique
                if field in SCIENTIFIC_FIELDS:

                    italicize_cell_species(cell)


                # couleurs statuts
                if field in STATUS_FIELDS:

                    apply_status_style_species(
                        cell,
                        str(value)
                    )

        # =====================================================
        # INSERTION
        # =====================================================

        paragraph._element.addnext(
            table._element
        )

        break