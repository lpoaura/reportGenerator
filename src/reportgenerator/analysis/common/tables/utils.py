from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt, RGBColor


from analysis.common.tables.styles import (
    LPO_BLUE,
    LPO_WHITE,
    LPO_GREY,
    STATUS_COLORS
)



def style_header_cell(cell):

    set_cell_background(cell, LPO_BLUE)

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = paragraph.runs[0]

    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(LPO_WHITE)


def apply_status_style(cell, value):

    if value in STATUS_COLORS:
        set_cell_background(cell, STATUS_COLORS[value])


def italicize_cell(cell):

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.italic = True


def set_cell_background(cell, color):

    shading_elm = parse_xml(
        rf'<w:shd {nsdecls("w")} w:fill="{color}"/>'
    )

    cell._tc.get_or_add_tcPr().append(
        shading_elm
    )


def set_cell_text_color(cell, color):

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)

bytes.fromhex
def set_cell_font(
    cell,
    bold=False,
    italic=False,
    size=10,
    font_name="LPO-Regular"
):

    for paragraph in cell.paragraphs:

        for run in paragraph.runs:

            run.font.bold = bold
            run.font.italic = italic
            run.font.size = Pt(size)
            run.font.name = font_name


def italicize_cell_species(cell):

    for paragraph in cell.paragraphs:

        for run in paragraph.runs:

            run.italic = True
            run.font.name = "LPO-Italic"
            run.font.size = Pt(10)
            run.center = True
            run.font.color.rgb = RGBColor.from_string(LPO_GREY)


def apply_status_style_species(cell, status):

    status = str(status).strip().upper()

    colors = {
        "EX": ("000000", "FFFFFF"),
        "EW": ("3D1851", "FFFFFF"),
        "RE": ("5B1A62", "FFFFFF"),
        "CR": ("E62328", "FFFFFF"),
        "EN": ("FFC33C", "000000"),
        "VU": ("FFED00", "000000"),
        "NT": ("FAF2C7", "000000"),
        "LC": ("78B747", "000000"),
        "DD": ("D4D4D4", "000000"),
    }

    if status not in colors:
        return

    bg_color, text_color = colors[status]

    set_cell_background(
        cell,
        bg_color
    )
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(text_color)