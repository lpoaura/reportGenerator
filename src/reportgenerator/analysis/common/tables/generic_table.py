from docx.shared import Inches
from analysis.common.tables.utils import style_header_cell


def insert_generic_table(
    document,
    placeholder,
    data,
    columns
):

    for paragraph in document.paragraphs:

        if placeholder in paragraph.text:

            paragraph.text = paragraph.text.replace(
                placeholder,
                ""
            )

            table = document.add_table(
                rows=1,
                cols=len(columns)
            )

            table.style = "Table Grid"

            # HEADER
            header_cells = table.rows[0].cells

            for idx, (_, label) in enumerate(columns):

                header_cells[idx].text = label
                style_header_cell(header_cells[idx])

            # DATA
            for item in data:

                row_cells = table.add_row().cells

                for idx, (field, _) in enumerate(columns):

                    value = item.get(field, "")
                    row_cells[idx].text = str(value)

            paragraph._element.addnext(table._element)

            break

def insert_general_table(document, placeholder, data, columns):

    for paragraph in document.paragraphs:

        if placeholder in paragraph.text:

            paragraph.text = paragraph.text.replace(placeholder, "")
            table = document.add_table(
                rows=1,
                cols=len(columns)
            )
            table.style = "Grid Table 4 Accent 6"
            # HEADER
            header = table.rows[0].cells

            for i, col in enumerate(columns):
                header[i].text = col[0]

            # DATA
            for item in data:
                #print("ITEM =", item)
                row = table.add_row().cells

                for i, col in enumerate(columns):
                    #print("COL =", col)
                    key = col[1]
                    #print("KEY =", key)
                    value = item.get(key, "")
                    #print("VALUE =", value)
                    row[i].text = str(value)
            paragraph._element.addnext(table._element)

            break