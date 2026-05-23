#!/bin/python3

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Mm
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


import pandas as pd
from datetime import datetime
from pathlib import Path

from db_auth import get_connection
from queries import SyntheseQueries
from analysis.common.tables.generic_table import insert_general_table, insert_generic_table
from analysis.common.tables.species_table import insert_species_table
from analysis.common.tables.excel_export import export_species_excel

LAYOUTS = {
    "normal": {
        "width": Inches(6),
        "height": None,
        "full_page": False,
    },

    "A4": {
        "width": Mm(170),
        "height": Mm(247),
        "full_page": True,
    },

    "A3": {
        "width": Mm(250),
        "height": Mm(370),
        "full_page": True,
    },
}

def replace_text(document, replacements: dict):

    # 1. Paragraphes classiques
    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            placeholder = "{{" + key + "}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    placeholder,
                    str(value)
                )

    # 2. Tableaux (IMPORTANT)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        placeholder = "{{" + key + "}}"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(
                                placeholder,
                                str(value)
                            )

## def insert_table_after_placeholder(document, placeholder, data):

    for paragraph in document.paragraphs:

        if placeholder in paragraph.text:

            paragraph.text = paragraph.text.replace(placeholder, "")

            table = document.add_table(rows=1, cols=10)

            header = table.rows[0].cells

            header[0].text = "Groupe taxonomique"
            header[1].text = "Nombre de données"
            header[2].text = "Nombre d'espèces"
            header[3].text = "Nombre d'espèces nicheuses"
            header[4].text = "Nombre d'espèces protégées"
            header[5].text = "Nombre d'espèces protégées nicheuses"
            header[6].text = "Nombre d'espèces en danger"
            header[7].text = "Nombre d'espèces en danger nicheuses"
            header[8].text = "Nombre de données mortalité"
            header[9].text = "Nombre espèces mortalité"

            for item in data:

                row = table.add_row().cells

                row[0].text = str(item["group_taxo"])
                row[1].text = str(item["nb_data_tot"])
                row[2].text = str(item["nb_espece"])
                row[3].text = str(item["nb_espece_nicheuse"])
                row[4].text = str(item["nb_espece_protege"])
                row[5].text = str(item["nb_espece_protege_nicheuse"])
                row[6].text = str(item["nb_espece_lr"])
                row[7].text = str(item["nb_espece_lr_nicheuse"])
                row[8].text = str(item["nb_data_mortalite"])
                row[9].text = str(item["nb_esp_mortalite"])

            paragraph._element.addnext(table._element)

            break
        
def insert_image(document, placeholder, image_path):
    for paragraph in document.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, "")
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(6))
            break

def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)

def insert_images_from_folder(document, folder: Path, layout="normal"):

    folder = Path(folder)
    config = LAYOUTS.get(layout)

    if not config:
        raise ValueError(f"Layout inconnu : {layout}")
    images = list(folder.glob("*.png"))

    for img in images:
        placeholder = f"{{{{{img.name}}}}}"
        for paragraph in document.paragraphs:
            if placeholder in paragraph.text:
                # suppression du placeholder
                paragraph.text = paragraph.text.replace(placeholder, "")
                # ---------- IMAGE NORMALE ----------
                if not config["full_page"]:
                    run = paragraph.add_run()
                    run.add_picture(
                        str(img),
                        width=config["width"]
                    )
                # ---------- IMAGE PLEINE PAGE ----------
                else:
                    # paragraphe après le placeholder
                    image_p = insert_paragraph_after(paragraph)
                    image_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = image_p.add_run()
                    run.add_picture(
                        str(img),
                        width=config["width"],
                        height=config["height"]
                    )
                    # saut de page après
                    page_break_p = insert_paragraph_after(image_p)
                    page_break_p.add_run().add_break()
                break

def generate_report(service_name: str, output_file: str, id_area: int, referee: str, list_analyse: str, buffer: int, area_name: str, analysis_result: dict):

    TEMPLATE_DIR = Path(__file__).resolve().parent / "templates"
    OUTPUT_DIR = Path(__file__).resolve().parent / "outputs" / area_name

    with get_connection(service_name) as conn:
        synthese_queries = SyntheseQueries(conn=conn, id_area=id_area, buffer=buffer)
        tableau_data = synthese_queries.get_resum_taxo_group()
        tableau_species = synthese_queries.get_species_data()
        tableau_resum = synthese_queries.get_resum_data()

    template_path = TEMPLATE_DIR / "Rapport_template.docx"
    dir_dataviz = OUTPUT_DIR / "dataviz"
    dir_maps = OUTPUT_DIR / "maps"
    dir_tables = OUTPUT_DIR / "tables"

    print(f"Génération du rapport Word à partir du template {output_file}...")

    document = Document(template_path)

    # 1 . Gestion des variables dans le fichiers word
    # ===============================================

    # récupération des données générales pour le résumé
    tableau_data_dict = tableau_data[0]
        
    zone_etude = next(
        row for row in tableau_resum
        if row["secteur"] == "zone_etude"
    )
    buffer_seul = next(
        row for row in tableau_resum
        if row["secteur"] == "buffer_seul"
    )
    ensemble = next(
        row for row in tableau_resum
        if row["secteur"] == "ensemble"
    )

    year_doc = '(' + str(int(datetime.now().strftime("%Y")) - 10) + '-' + datetime.now().strftime("%Y") + ')'

    replace_text(document, {
        "AREA_NAME": area_name,
        "REFEREE": referee,
        "DATE_REPORT": year_doc,
        "BUFFER": buffer,
        "NB_DATA": tableau_data_dict["nb_data_tot"],
        "NB_OBS_ZONE":f"{zone_etude['nb_observations']:,}".replace(",", " "),
        "NB_SPECIES_ZONE":f"{zone_etude['nb_especes']:,}".replace(",", " "),
        "LAST_OBS_ZONE":str(zone_etude["derniere_observation"]),
        "NB_OBS_BUFFER":f"{buffer_seul['nb_observations']:,}".replace(",", " "),
        "NB_SPECIES_BUFFER":f"{buffer_seul['nb_especes']:,}".replace(",", " "),
        "LAST_OBS_BUFFER":str(buffer_seul["derniere_observation"]),
        "NB_OBS_GLOBAL":f"{ensemble['nb_observations']:,}".replace(",", " "),
        "NB_SPECIES_GLOBAL":f"{ensemble['nb_especes']:,}".replace(",", " "),
        "LAST_OBS_GLOBAL":str(ensemble["derniere_observation"]),
    })

    # 2 . Gestion des tableaux dans le fichiers word
    # ===============================================
    # préparation des données pour les tableaux
    columns_taxo = [
        ('Groupe taxo.', 'group_taxo'),
        ('Nombre\nde données', 'nb_data_tot'),
        ("Nombre\nd'espèces", 'nb_espece'),
        ("Espèces\nnicheuses", 'nb_espece_nicheuse'),
        ("Espèces\nprotégées", 'nb_espece_protege'),
        ("Espèces\nen danger", 'nb_espece_lr'),
       # ("Espèces en danger nicheuses", 'nb_espece_lr_nicheuse'),
    ]


    STATUS_COLS = ["lr_auv", "lr_ra", "lr_aura", "lr_france"]
    STATUS_FILTER = ["CR", "EN", "VU", "NT"]

    tableau_species_lr = pd.DataFrame(tableau_species)
    mask = tableau_species_lr[STATUS_COLS].isin(STATUS_FILTER).any(axis=1)
    tableau_species_lr = tableau_species_lr[mask]
    tableau_species_lr = tableau_species_lr.to_dict( orient="records")
       
    # insertion des tableaux dans le document
    insert_species_table(
        document=document,
        placeholder="{{TABLE_ESP}}",
        data=tableau_species
    )

    insert_general_table(
        document=document,
        placeholder="{{TABLE_TAXO}}",
        data=tableau_data,
        columns=columns_taxo
    )   

    insert_species_table(
        document=document,
        placeholder="{{TABLE_ESP_LR}}",
        data=tableau_species_lr
    )


    # export Excel des tableaux 
    export_species_excel( data=tableau_species, output_path=(dir_tables / "tableau_especes.xlsx"))


    # 3 . Gestion des images dans le fichier word
    # ===============================================

    insert_images_from_folder(document, dir_dataviz, "normal")
    insert_images_from_folder(document, dir_maps, "A4")
    
    document.save(output_file)




